#!/usr/bin/env python3

"""
CV Generator

"""
from jinja2 import Environment, PackageLoader, select_autoescape
import ruamel.yaml as yaml
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from spellchecker import SpellChecker
import re
import numbers

import datetime

output_options = {
    # Should company given titles be shown?
    # For some applications it will be more suitable to define the role I was
    # doing in the summary text, than the company given title
    'show_position_titles': False,
    'show_skill_headings': False,
    'include_tags': ['base']
}
cvdb = {}

with open("cvdb.yml", 'r') as stream:
    try:
        cvdb = yaml.load(stream, Loader=yaml.Loader)
    except yaml.YAMLError as exc:
        print(exc)

# Run spell check across the database
spell = SpellChecker()
spell.word_frequency.load_words([
    "2G",
    "3G",
    "4-year",
    "8am",
    "8pm",
    "9FU",
    "ABR",
    "AVG",
    "AWS",
    "AmiNET",
    "Amino",
    "AminoMOVE",
    "Ansible",
    "Aquent",
    "Berryville",
    "Booxmedia",
    "Bugzilla",
    "Buildroot",
    "CB25",
    "CMake",
    "CircleCI",
    "CloudTV",
    "Cython",
    "Debian",
    "DockerHub",
    "Entone",
    "Facebook",
    "Fachhochschule",
    "GPRS",
    "GitHub",
    "Google",
    "HLS",
    "HTML",
    "HTML5",
    "IGMP",
    "ISEB",
    "JIRA",
    "Javascript",
    "Landbeach",
    "Linux",
    "LogDevice",
    "MP3",
    "MSTV",
    "MySQL",
    "PHP",
    "PRINCE2",
    "Perl",
    "Rsyslog",
    "STB",
    "Splunk",
    "TTPCom",
    "TravisCI",
    "TestTrack",
    "Unomaly",
    "WiMAX",
    "Yii",
    "Zabbix",
    "Autotools",
    "caleb.marchent@icloud.com",
    "fbthrift-py3",
    "preemptively"
    ])

def check_text_spelling(text):
    try:
        # Handle hyphenated or underscored words, by replacing with spaces
        text = text.replace("_", " ")
        words = text.split()
        # Strip out punctuation and bold text markers
        words = [ word.rstrip(".,;:)*").lstrip("*($") for word in words ]
        # Strip off possessive apostrophe-s
        words = [ word[:-2] if word.endswith("'s")
                            or word.endswith("’s") else word for word in words ]
        # Drop any items that resulted in blank spaces
        words = [ word for word in words if len(word) > 0 ]
    except:
        raise ValueError("Failed to split:\n \"{}\"".format(text))
    unknown = spell.unknown(words)
    # Split any hypenated words and check again, this would ensure "known
    # hypenated words" are handled correctly, before falling back
    dehyphenated = [ word for c_list in unknown for word in c_list.split("-") ]
    bad = spell.unknown(dehyphenated)
    for word in bad:
        print ("Spelling error {}".format(word))


def spellcheck_struct(s):
    if isinstance(s, dict):
        for k, v in s.items():
            spellcheck_struct(v)
    elif isinstance(s, list):
        for item in s:
            spellcheck_struct(item)
    elif isinstance(s, str):
        check_text_spelling(s)
    elif s is not None and not isinstance(s, numbers.Number):
        raise ValueError("Unexpected item in the spellcheck_struct : \"{}\"".format(s))

spellcheck_struct(cvdb)

# Process the imported YAML; creating the structure required for the document
# from the source database
processed_cvdb = dict()
processed_cvdb['skill_groups'] = list()
row = 0
for grp in cvdb['skill_groups']:
    d = dict([('heading', grp['heading']), ('items', grp['items'])])
    processed_cvdb['skill_groups'].append(d)
#    for skill in grp['items']:
#        if 'hidden' not in cvdb['skill_groups']:
#            processed_cvdb['skill_groups'][row]['items'].append(skill)
    row += 1

processed_cvdb['contact_info'] = " | ".join((
    cvdb['contact_info']['postal_address'],
    cvdb['contact_info']['phone_number'],
    cvdb['contact_info']['email_address']))

processed_cvdb['positions'] = cvdb['positions']
processed_cvdb['elevator_pitch'] = cvdb['elevator_pitch']
processed_cvdb['further_information'] = cvdb['further_information']

processed_cvdb['achievements'] = {}
for position, position_achievements in cvdb['achievements'].items():
    processed_cvdb['achievements'][position] = list()
    for achievement in position_achievements:
        if ('tags' not in achievement) or bool(set(output_options['include_tags']) & set(achievement['tags'])):
            processed_cvdb['achievements'][position].append(achievement)

processed_cvdb['education'] = cvdb['education']

OUTPUT_DOCUMENT_TYPE = "word"  # or "html"

env = Environment(
    loader=PackageLoader('gencv', 'templates'),
    autoescape=select_autoescape(['html', 'rtf'])
)

# Update the media wiki version of the document
template = env.get_template('cv_mediawiki.md')

# to save the results
with open("README.md", "wb") as fh:
    fh.write(template.render(processed_cvdb).encode('utf-8'))

if OUTPUT_DOCUMENT_TYPE == "html":


    template = env.get_template('legacy_cv.html')

    # to save the results
    with open("generated_cv.html", "wb") as fh:
        fh.write(template.render(processed_cvdb).encode('utf-8'))

elif OUTPUT_DOCUMENT_TYPE == "word":

    document = Document()

#   List known styles and types; useful for finding correct style name
    # for style in document.styles: print "{} ({})".format(style.name, style.type)

    for section in document.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.page_height = Mm(297)
        section.page_width = Mm(210)

    styles = document.styles

    styles['Title'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles['Title'].paragraph_format.space_after = Pt(3)
    styles['Normal'].font.size = Pt(10)

    s = styles.add_style('skills_bullet', WD_STYLE_TYPE.PARAGRAPH)
    s.base_style = styles['List Bullet']
    s.font.size = Pt(10)

    s = styles.add_style('achievement_bullet', WD_STYLE_TYPE.PARAGRAPH)
    s.base_style = styles['List Bullet']
    s.font.size = Pt(10)

    s = styles.add_style('company_summary', WD_STYLE_TYPE.PARAGRAPH)
    s.base_style = styles['Normal']
    s.font.italic = True

    s = styles.add_style('contact_info', WD_STYLE_TYPE.PARAGRAPH)
    s.base_style = styles['Normal']
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.font.size = Pt(9)

    # Configure the document properties:
    document.core_properties.title = "Caleb Marchent"
    document.core_properties.author = "Caleb Marchent - Curriculum Vitae"
    document.core_properties.created = datetime.datetime.now()
    document.core_properties.comments = \
        "Generated Automatically\nSource code available at:\n" + \
        "https://github.com/calebmarchent/auto_cv.git"

    document.add_heading('Caleb Marchent', 0)

    p = document.add_paragraph(processed_cvdb['contact_info'],
                               style='contact_info')

    document.add_heading('Profile', level=2)

    for para in processed_cvdb['elevator_pitch']:
        p = document.add_paragraph(para)

    document.add_heading('Key Skills', level=2)

    # Find the maximum number of skills in a group
    table = document.add_table(rows=(2 if output_options['show_skill_headings'] else 1),
                               cols=len(processed_cvdb['skill_groups']))
    col = 0

    # Populate table with skills, each cell will have a default paragraph
    # generated automatically by the API, as not having one is invalid syntax;
    # for the first paragraph we need to update where we 'add' for the other
    # bullets

    for skill_group in processed_cvdb['skill_groups']:
        row = 0
        idx = 0
        hdr_cells = table.columns[col].cells
        if output_options['show_skill_headings']:
            hdr_cells[row].paragraphs[0].add_run(skill_group['heading']).bold = True
            row += 1
        for skill in skill_group['items']:
            if idx == 0:
                hdr_cells[row].paragraphs[0].style = styles['skills_bullet']
                p = hdr_cells[row].paragraphs[0]
            else:
                p = hdr_cells[row].add_paragraph('', style='skills_bullet')

            # Look for embolden sections and split run to illuminate the text
            # between them
            # FIXME: The code below alternates boldness on each occurance of *,
            # while it would work for now, it is not extensible to handle other
            # formatting charaters.
            b = False
            for run in skill.split('*'):
                p.add_run(run).bold = b
                b = not b
            idx += 1

        col += 1

    document.add_heading('Experience', level=2)

    for position in processed_cvdb['positions']:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)

        p.add_run(position['company_name']).bold = True
        if output_options['show_position_titles']:
            p.add_run("\t" + position['title'] + "\t")
        else:
            p.add_run("\t\t")

        if 'finish' in position:
           finish_txt = position['finish']
        else:
           finish_txt = 'to date'
        p.add_run(str(position['start']) + " - " + str(finish_txt)).bold = True

        p.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.CENTER)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)

        if 'company_summary' in position:
            document.add_paragraph(position['company_summary'], style='company_summary')

        if position['brief'] in processed_cvdb['achievements']:
            for achievement in processed_cvdb['achievements'][position['brief']]:
                p = document.add_paragraph(achievement['desc'],
                    style='achievement_bullet')

    document.add_heading('Education', level=2)

    for experience in processed_cvdb['education']:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run("{}\t{}".format(experience['desc'], experience['date'])).bold = True
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
        if 'additional_info' in experience:
            p.add_run("\n{}".format(experience['additional_info']))

    document.add_heading('Further Information', level=2)

    for line in processed_cvdb['further_information']:
        p = document.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    document.save('curriculum_vitae.docx')
